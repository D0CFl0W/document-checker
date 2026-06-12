from __future__ import annotations

import os
import platform
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

import cv2
import fitz
import numpy as np
import pytesseract
from docx import Document
from PIL import Image
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

BASE_DIR = Path(__file__).resolve().parent.parent
FONT_PATH = BASE_DIR / "fonts" / "DejaVuSans.ttf"


def _detect_format(document_path: Path) -> str:
    try:
        head = document_path.read_bytes()[:8]
    except OSError as exc:
        raise ValueError(f"Не удалось прочитать файл {document_path}: {exc}") from exc

    if head.startswith(b"%PDF"):
        return ".pdf"
    if head.startswith(b"PK\x03\x04") or head.startswith(b"PK\x05\x06"):
        return ".docx"

    suffix = document_path.suffix.lower()
    if suffix == ".pdf":
        return ".pdf"
    if suffix == ".docx":
        return ".docx"
    raise ValueError(
        f"Неподдерживаемый формат (ожидались PDF или DOCX): "
        f"{document_path.suffix or '(без расширения)'}",
    )


def extract_pdf_ocr(document_path: Path, lang: str = "rus+eng", zoom: float = 2.0) -> str:
    try:
        doc = fitz.open(str(document_path.resolve()))
        text_parts: list[str] = []

        mat = fitz.Matrix(zoom, zoom)
        config = r"--oem 3 --psm 6"

        for page in doc:
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

            img = img.convert("L")

            page_text = pytesseract.image_to_string(
                img,
                lang=lang,
                config=config
            )

            text_parts.append(page_text)

        return "\n".join(text_parts)

    except Exception as exc:
        raise ValueError(f"Ошибка OCR PDF {document_path}: {exc}") from exc


def extract_docx(document_path: Path) -> str:
    try:
        doc = Document(document_path)
    except Exception as exc:
        raise ValueError(f"Ошибка чтения файла {document_path}: {exc}") from exc
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def extract_text(path: str | Path) -> str:
    document_path = Path(path)
    kind = _detect_format(document_path)

    if kind == ".docx":
        return extract_docx(document_path)

    return extract_pdf_ocr(document_path)


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[|•·]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n+", "\n", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def fix_ocr_errors(text: str) -> str:
    replacements = {
        "Гpyппа": "Группа",
        "Групnа": "Группа",
        "TeMa": "Тема",
        "Teмa": "Тема",
        "CTyдeнт": "Студент",
        "Cтудент": "Студент",
        "ФИ0": "ФИО",
        "oбучающегося": "обучающегося",
        "»": "",
    }

    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)

    return text


patterns: dict[str, list[str]] = {
    "ФИО": [
        r"(?:Студент|студенту|Исполнитель|обучающегося)\s*[:\-]?\s*(?P<value>[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.)",
    ],
    "Группа": [
        r"(?:Группа|Группы)\s*[:\-]?\s*(?P<value>[A-ЯA-ZЁ]{1,4}\s*-?\s*\d{2,6})",
        r"(?:группа)\s*(?P<value>[A-ЯA-ZЁ]{1,4}\s*-?\s*\d{2,6})",
    ],
    "Тема": [
        r"(?:Тема работы:|Тема ВКР|Тема|на тему)\s*[:\-]?\s*«?(?P<value>[^_\n»]{5,200})",
    ],
    "Дата": [
        r"Дата:\s*(?P<value>\d{2}\.\d{2}\.\d{2,4}|\[.*?\]|_+)",
        r"Срок сдачи студентом.*?:\s*(?P<value>\d{2}\.\d{2}\.\d{2,4}|\[.*?\]|_+)",
        r"(?:Дата|Срок сдачи|Дата проверки|Работа сдана)\s*[:\-]?\s*"
        r"«?\s*(?P<value>\d{1,2}\s*[.\-]\s*\d{1,2}\s*[.\-]\s*\d{2,4})\s*»?",
        r"(?P<value>\d{1,2}\s+[а-яА-ЯёЁ]{3,10}\s+\d{4})",
        r"(?P<value>\d{1,2}\s*[а-яА-ЯёЁ]{3,10}\s*\d{4}\s*[гr]?)",
        r"«\s*(?P<value>\d{1,2}\s*[а-яА-ЯёЁ]+\s*\d{4})\s*»"
    ],
    "Подпись": [],
}


def clean_fio(value: str | None) -> str | None:
    if not value:
        return None

    value = re.sub(r"\s+", " ", value).strip()

    if re.search(r"\b(групп|посвящен|работ|тема|дата)\b", value.lower()):
        return None

    value = re.sub(r"\.\s*", ". ", value)
    value = re.sub(r"\s+", " ", value).strip()

    return value


def clean_value(value: str) -> str | None:
    stripped = value.strip()
    if re.fullmatch(r"_+", stripped):
        return None
    if re.fullmatch(r"\[.*?\]", stripped):
        return None
    without_trailing_paren = re.sub(r"\s*\(.*?\)\s*$", "", stripped)
    return without_trailing_paren.lstrip(":").strip()


def extract_field(text: str, field_patterns: list[str]) -> str | None:
    for pattern in field_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE | re.DOTALL):
            value = clean_value(match.group("value"))

            if not value:
                continue

            if re.search(r"\b(группа|тема|дата|работа|курсовая)\b", value.lower()):
                continue

            return value

    return None


def extract_topic(text: str) -> str | None:
    match = re.search(
        r"(?:Тема|Тема работы|Тема ВКР|на тему)\s*[:\-]?\s*«?(?P<value>[^\n»]+)",
        text,
        re.IGNORECASE
    )

    if not match:
        return None

    value = match.group("value").strip()

    if "_" in value:
        return None

    return value


def clean_topic(value: str | None) -> str | None:
    if not value:
        return None

    value = value.strip()
    if "_" in value:
        return None

    bad_fragments = [
        "антиплагиат",
        "система проверки",
        "удк",
        "институт",
        "рег.",
        "инв.",
    ]

    low = value.lower()
    if any(b in low for b in bad_fragments):
        return None

    cut_markers = ["система проверки", "удк", "рег.", "инв.", "институт"]

    for m in cut_markers:
        if m in low:
            value = re.split(m, value, flags=re.IGNORECASE)[0].strip()

    return value if value else None


def _fallback_docx_check(document_path: Path) -> str:
    try:
        doc = Document(document_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)

        if re.search(r"_{5,}", all_text) or re.search(r"-{5,}", all_text):
            return "нет"

        for i, para in enumerate(doc.paragraphs):
            if "подпись" in para.text.lower():
                if len(para.inline_shapes) > 0:
                    return "есть"
                if i + 1 < len(doc.paragraphs) and len(doc.paragraphs[i + 1].inline_shapes) > 0:
                    return "есть"
        return "нет"
    except Exception:
        return "нет"


def detect_signature_cv2(document_path: Path, kind: str, debug: bool = False) -> str:
    actual_path = document_path
    temp_dir = None

    try:
        if kind == ".docx":
            if platform.system() == "Windows":
                win_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                cmd_exec = None
                for p in win_paths:
                    if os.path.exists(p):
                        cmd_exec = p
                        break

                if not cmd_exec:
                    return _fallback_docx_check(document_path)
            else:
                cmd_exec = "libreoffice"

            temp_dir = Path(tempfile.mkdtemp())
            env = os.environ.copy()
            env["HOME"] = str(temp_dir)
            lo_profile_path = (temp_dir / "lo_profile").as_uri()

            cmd = [
                cmd_exec, "--headless", "--invisible", "--norestore",
                f"-env:UserInstallation={lo_profile_path}",
                "--convert-to", "pdf",
                "--outdir", str(temp_dir),
                str(document_path.resolve())
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30, env=env)

            if result.returncode != 0:
                err_msg = result.stderr.strip().replace('\n', ' ')[:150]
                return f"нет (ошибка конвертации: {err_msg})"

            pdf_files = list(temp_dir.glob("*.pdf"))
            if not pdf_files:
                return "нет"

            actual_path = pdf_files[0]
            kind = ".pdf"

        doc = fitz.open(str(actual_path.resolve()))
        pages_to_check = doc[-2:] if len(doc) >= 2 else doc

        for page_idx, page in enumerate(pages_to_check):
            hits = page.search_for("Подпись") or page.search_for("подпись")

            if hits:
                rect = hits[0]
                crop_y0 = max(0, rect.y0 - 100)
                crop_y1 = min(page.rect.height, rect.y1 + 50)
                crop_x0 = max(0, rect.x0 - 50)
                crop_x1 = min(page.rect.width, rect.x1 + 200)
            else:
                crop_y0 = int(page.rect.height * 0.65)
                crop_y1 = int(page.rect.height)
                crop_x0 = 0
                crop_x1 = int(page.rect.width)

            crop_rect = fitz.Rect(crop_x0, crop_y0, crop_x1, crop_y1)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, clip=crop_rect)

            if pix.width < 50 or pix.height < 50:
                continue

            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

            thresh = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY_INV, 15, 10
            )

            if debug:
                cv2.imwrite(f"debug_page{page_idx}_crop.png", img)
                cv2.imwrite(f"debug_page{page_idx}_thresh.png", thresh)

            h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (60, 1))
            horizontal_lines = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, h_kernel)

            if debug:
                cv2.imwrite(f"debug_page{page_idx}_lines.png", horizontal_lines)

            contours, _ = cv2.findContours(horizontal_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

            signature_found = False

            for cnt in contours:
                x, y, w, h = cv2.boundingRect(cnt)

                if w > 50 and h < 15:
                    roi_y_start = max(0, y - 80)
                    roi_y_end = max(roi_y_start + 10, y - 5)

                    roi_x_start = max(0, x - 30)
                    roi_x_end = min(thresh.shape[1], x + w + 30)

                    roi = thresh[roi_y_start:roi_y_end, roi_x_start:roi_x_end]

                    ink_pixels = cv2.countNonZero(roi)

                    if ink_pixels > 800:
                        signature_found = True
                        break

            if signature_found:
                return "есть"

        return "нет"

    except subprocess.TimeoutExpired:
        return "нет (таймаут)"
    except Exception as exc:
        return f"ошибка: {exc}"
    finally:
        if temp_dir and temp_dir.exists():
            shutil.rmtree(temp_dir, ignore_errors=True)


def parse_document(path: str | Path) -> dict[str, str | None]:
    document_path = Path(path)
    kind = _detect_format(document_path)

    raw_text = extract_text(document_path)
    normalized = normalize_text(raw_text)

    result: dict[str, str | None] = {}

    for field, field_patterns in patterns.items():
        if field == "Подпись":
            result[field] = detect_signature_cv2(document_path, kind)
            continue

        if field == "Тема":
            topic = extract_field(normalized, field_patterns) or extract_topic(raw_text)
            result[field] = clean_topic(topic)
            continue

        value = extract_field(normalized, field_patterns)

        if field == "ФИО":
            value = clean_fio(value)

        result[field] = value

    return result


def parse_directory(directory: str | Path, limit: int = 5) -> list[dict[str, Any]]:
    root = Path(directory)
    if not root.exists():
        raise ValueError(f"Директория не существует: {root}")

    docx_files = list(root.glob("*.docx"))
    pdf_files = list(root.glob("*.pdf"))
    files_sorted = sorted(docx_files + pdf_files)
    if limit:
        files_sorted = files_sorted[:limit]

    results: list[dict[str, Any]] = []
    for file_path in files_sorted:
        try:
            data = parse_document(file_path)
            results.append({"file": file_path.name, "data": data})
        except Exception as exc:
            results.append({"file": file_path.name, "error": str(exc)})
    return results


KEY_FIELDS = ["ФИО", "Группа", "Тема"]


def _normalize_fio(value: str) -> str | None:
    words = re.findall(r'[а-яё]+', value.lower())
    if not words:
        return None
    words.sort(key=len, reverse=True)
    return words[0][:4] if len(words[0]) >= 4 else words[0]


def _normalize_group(value: str) -> str | None:
    return re.sub(r'[^а-яёa-z0-9]', '', value.lower())


def _normalize_topic(value: str) -> str | None:
    clean_topic = re.sub(r'[^\w\s]', '', value.lower())
    words = clean_topic.split()
    return " ".join(words[:3])


def evaluate_completeness(results: list[dict[str, Any]]) -> tuple[bool, dict[str, set[str]]]:
    inconsistencies: dict[str, set[str]] = {}

    for field in KEY_FIELDS:
        normalized_values = set()
        original_values = set()

        for item in results:
            if "error" in item:
                continue

            raw_val = item["data"].get(field)

            if not raw_val or str(raw_val).strip() in ["—", "-", "нет", ""]:
                continue

            if field == "ФИО":
                norm_val = _normalize_fio(str(raw_val))
            elif field == "Группа":
                norm_val = _normalize_group(str(raw_val))
            elif field == "Тема":
                norm_val = _normalize_topic(str(raw_val))
            else:
                norm_val = str(raw_val).strip().lower()

            if norm_val:
                normalized_values.add(norm_val)
                original_values.add(str(raw_val).strip())

        if len(normalized_values) > 1:
            inconsistencies[field] = original_values

    is_complete = len(inconsistencies) == 0
    return is_complete, inconsistencies



def generate_pdf_report(results: list[dict[str, Any]], output_path: str) -> None:
    pdf_canvas = canvas.Canvas(output_path)
    pdfmetrics.registerFont(TTFont("DejaVu", str(FONT_PATH)))
    pdf_canvas.setFont("DejaVu", 9)

    is_complete, inconsistencies = evaluate_completeness(results)

    y_position = 800

    status = "КОМПЛЕКТ" if is_complete else "НЕ КОМПЛЕКТ"
    pdf_canvas.drawString(50, y_position, f"Статус набора: {status}")
    y_position -= 15

    if inconsistencies:
        pdf_canvas.drawString(50, y_position, "Несоответствия:")
        y_position -= 10

        for field, values in inconsistencies.items():
            readable = [v if v is not None else "—" for v in values]
            pdf_canvas.drawString(60, y_position, f"{field}: {', '.join(readable)}")
            y_position -= 10

        y_position -= 10

    for item in results:
        if y_position < 100:
            pdf_canvas.showPage()
            pdf_canvas.setFont("DejaVu", 9)
            y_position = 800

        if "error" in item:
            pdf_canvas.drawString(
                50, y_position, f"{item['file']} — ОШИБКА: {item['error']}"
            )
            y_position -= 10
            continue

        pdf_canvas.drawString(10, y_position, f"Файл: {item['file']}")
        y_position -= 10

        for key, value in item["data"].items():
            display = value if value else "—"
            pdf_canvas.drawString(20, y_position, f"{key}: {display}")
            y_position -= 10

        y_position -= 8

    pdf_canvas.save()
